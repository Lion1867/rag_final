import re
import os
from typing import List, Dict


MIME_TYPES = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}


def get_file_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    return ext.lstrip(".")


def is_supported(file_path: str) -> bool:
    ext = os.path.splitext(file_path)[1].lower()
    return ext in SUPPORTED_EXTENSIONS


def get_doc_name(file_path: str) -> str:
    basename = os.path.basename(file_path)
    name, _ = os.path.splitext(basename)
    return name


def extract_text_from_pdf_with_pages(file_path: str) -> List[Dict]:
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF ne ustanovlen: pip install PyMuPDF")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = fitz.open(file_path)
    pages = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages.append({
                "text": text,
                "page": page_num + 1
            })

    doc.close()
    print(f"   PDF: {len(pages)} pages")
    return pages


def extract_text_from_docx_with_pages(file_path: str) -> List[Dict]:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx ne ustanovlen: pip install python-docx")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))

    full_text = "\n\n".join(paragraphs)
    print(f"   DOCX: {len(paragraphs)} paragraphs")
    return [{"text": full_text, "page": 1}]


def extract_text_from_doc_with_pages(file_path: str) -> List[Dict]:
    import subprocess

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            print(f"   DOC: antiword")
            return [{"text": result.stdout, "page": 1}]
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    try:
        result = subprocess.run(
            ["catdoc", "-w", file_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            print(f"   DOC: catdoc")
            return [{"text": result.stdout, "page": 1}]
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    try:
        return extract_text_from_docx_with_pages(file_path)
    except:
        pass

    raise RuntimeError("Cannot extract DOC. Install antiword: apt-get install antiword")


def extract_pages(file_path: str) -> List[Dict]:
    file_type = get_file_type(file_path)
    if file_type == "pdf":
        return extract_text_from_pdf_with_pages(file_path)
    elif file_type == "docx":
        return extract_text_from_docx_with_pages(file_path)
    elif file_type == "doc":
        return extract_text_from_doc_with_pages(file_path)
    else:
        raise ValueError(f"Unsupported format: {file_path}")


def clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\n +", "\n", text)
    text = re.sub(r"[\x00-\x09\x0b-\x0c\x0e-\x1f]", "", text)
    return text.strip()


def split_page_text_into_chunks(
    text: str,
    page: int,
    chunk_size: int = 800,
    chunk_overlap: int = 200,
    start_chunk_id: int = 0,
) -> List[Dict]:
    text = clean_text(text)
    if not text:
        return []

    if len(text) <= chunk_size:
        return [{"text": text, "chunk_id": start_chunk_id, "page": page}]

    chunks = []
    start = 0
    chunk_id = start_chunk_id

    while start < len(text):
        end = min(start + chunk_size, len(text))

        if end < len(text):
            search_from = start + int(chunk_size * 0.7)
            best_break = -1

            for sep in [".\n", ". ", "!\n", "! ", "?\n", "? ", ";\n", "; ", "\n\n", "\n"]:
                pos = text.rfind(sep, search_from, end)
                if pos != -1 and pos + len(sep) > best_break:
                    best_break = pos + len(sep)

            if best_break > search_from:
                end = best_break

        chunk_text = text[start:end].strip()

        if chunk_text and len(chunk_text) >= 50:
            chunks.append({"text": chunk_text, "chunk_id": chunk_id, "page": page})
            chunk_id += 1

        next_start = end - chunk_overlap
        if next_start <= start:
            next_start = end
        start = next_start

    return chunks


def process_document(
    file_path: str,
    chunk_size: int = 800,
    chunk_overlap: int = 200,
    use_api: bool = False,
) -> List[Dict]:
    filename = os.path.basename(file_path)
    doc_name = get_doc_name(file_path)
    file_type = get_file_type(file_path)

    print(f"Processing: {filename} [{file_type.upper()}]")

    if not is_supported(file_path):
        print(f"   Unsupported format")
        return []

    try:
        pages = extract_pages(file_path)
    except Exception as e:
        print(f"   Error: {e}")
        return []

    total_chars = sum(len(p["text"]) for p in pages)
    print(f"   Extracted chars: {total_chars}")

    if total_chars == 0:
        print("   Document has no text")
        return []

    all_chunks = []
    chunk_id_counter = 0

    for page_info in pages:
        page_chunks = split_page_text_into_chunks(
            page_info["text"],
            page_info["page"],
            chunk_size,
            chunk_overlap,
            start_chunk_id=chunk_id_counter,
        )
        for ch in page_chunks:
            ch["source"] = doc_name
            ch["filename"] = filename
        chunk_id_counter += len(page_chunks)
        all_chunks.extend(page_chunks)

    print(f"   Created chunks: {len(all_chunks)}")

    if all_chunks:
        preview = all_chunks[0]["text"][:120].replace("\n", " ")
        print(f"   Preview: {preview}...")

    return all_chunks


def process_pdf(file_path: str, chunk_size: int = 800, chunk_overlap: int = 200, use_api: bool = False) -> List[Dict]:
    return process_document(file_path, chunk_size, chunk_overlap, use_api)